"""
Research Module for AmaniQuery
Uses Gemini AI to analyze legal queries and generate reports on Kenya's laws
"""

import os
import json
from typing import Dict, List, Optional, Any
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

class ResearchModule:
    """
    Research module using Gemini AI for analyzing legal queries
    and generating reports on Kenya's laws and legal information
    """

    def __init__(self, api_key: Optional[str] = None):
        """
        Initialize research module with Gemini API

        Args:
            api_key: Gemini API key (optional, will use env var if not provided)
        """
        self.api_key = api_key or os.getenv("GEMINI_API_KEY")
        if not self.api_key:
            raise ValueError("GEMINI_API_KEY not set in environment")

        try:
            import google.generativeai as genai
            genai.configure(api_key=self.api_key)
            self.genai = genai  # Store reference for later use
            self.model = genai.GenerativeModel('gemini-2.5-flash')
            logger.info("Research module initialized with Gemini AI")
        except ImportError:
            raise ValueError("google-generativeai package not installed. Install with: pip install google-generativeai")

    def analyze_legal_query(self, query: str, context: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Analyze a legal query about Kenya's laws with comprehensive research

        Args:
            query: The legal question or query to analyze
            context: Additional context about the query (optional)

        Returns:
            Dictionary containing detailed analysis results
        """
        context_info = ""
        if context:
            context_info = f"\n\nADDITIONAL CONTEXT:\n{json.dumps(context, indent=2)}"

        prompt = f"""
You are a senior legal researcher specializing in Kenyan law. Conduct comprehensive legal research and analysis for the following query:

LEGAL QUERY: {query}{context_info}

Provide a detailed research report in the following structured JSON format:

{{
  "original_query": "{query}",
  "executive_summary": "Concise overview of key findings, legal implications, and main recommendations (2-3 paragraphs)",
  "background_context": "Historical, social, and legal context relevant to the query (comprehensive background)",
  "methodology": "Research methodology, sources consulted, and analytical approach used",
  "applicable_laws": [
    {{
      "law_name": "Full name of the law/act",
      "citation": "Official citation (e.g., Cap 123 Laws of Kenya)",
      "key_provisions": ["Specific sections relevant to the query"],
      "amendments": "Recent amendments or changes",
      "interpretation": "How the law applies to this query"
    }}
  ],
  "constitutional_analysis": {{
    "relevant_articles": ["Specific constitutional articles"],
    "fundamental_rights": ["Rights implicated"],
    "constitutional_remedies": ["Available constitutional remedies"],
    "court_jurisdiction": "Which courts have jurisdiction"
  }},
  "case_law_precedents": [
    {{
      "case_name": "Full case citation",
      "court": "Court that decided the case",
      "year": "Year decided",
      "key_holding": "Main legal principle established",
      "relevance": "How it applies to this query"
    }}
  ],
  "detailed_legal_analysis": "Comprehensive legal analysis including step-by-step reasoning, legal principles, and implications (detailed explanation)",
  "practical_guidance": {{
    "immediate_steps": ["Step-by-step actions to take"],
    "required_documents": ["Documents needed"],
    "relevant_institutions": ["Government agencies, courts, or organizations involved"],
    "procedural_requirements": "Detailed procedures to follow",
    "timeframes": "Relevant deadlines and time considerations",
    "costs": "Associated costs and fees"
  }},
  "risk_assessment": {{
    "legal_risks": ["Potential legal risks and consequences"],
    "compliance_requirements": ["Legal obligations to fulfill"],
    "mitigation_strategies": ["Ways to minimize risks"],
    "alternative_approaches": ["Alternative legal strategies"]
  }},
  "recommendations": [
    {{
      "priority": "High/Medium/Low",
      "action": "Specific recommended action",
      "rationale": "Reason for the recommendation",
      "timeline": "When to implement",
      "responsible_party": "Who should take action"
    }}
  ],
  "additional_considerations": {{
    "related_legal_areas": ["Other areas of law that may be relevant"],
    "ethical_considerations": "Ethical issues to consider",
    "policy_implications": "Broader policy or societal implications",
    "future_developments": "Upcoming legal changes or developments"
  }},
  "sources_and_references": [
    {{
      "type": "Primary/Secondary",
      "title": "Source title",
      "author": "Author or institution",
      "publication_date": "Date published",
      "url": "Web link if available",
      "relevance": "How it was used in this analysis"
    }}
  ],
  "disclaimer": "Comprehensive legal disclaimer emphasizing that this is not formal legal advice and professional consultation is recommended",
  "model_used": "gemini-research",
  "research_timestamp": "{datetime.utcnow().isoformat()}",
  "report_confidence": "High/Medium/Low confidence in the analysis based on available information"
}}

Ensure the analysis is comprehensive, accurate, and based on Kenyan legal framework. Include specific references to laws, cases, and legal principles. Structure the response as valid JSON with detailed, actionable information.
"""

        try:
            response = self.model.generate_content(
                prompt,
                generation_config=self.genai.types.GenerationConfig(
                    temperature=0.2,
                    max_output_tokens=12000,
                    response_mime_type="application/json"
                )
            )
            
            # Parse the JSON response
            result = json.loads(response.text.strip())
            
            # Ensure required fields are present
            if 'research_timestamp' not in result:
                result['research_timestamp'] = datetime.utcnow().isoformat()
            
            if 'model_used' not in result:
                result['model_used'] = 'gemini-research'
                
            return result
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error in legal analysis: {e}")
            # Fallback to text processing
            return self._fallback_text_analysis(query, context)
        except Exception as e:
            logger.error(f"Error in legal query analysis: {e}")
            return {
                "original_query": query,
                "error": f"Research analysis failed: {str(e)}",
                "executive_summary": "Unable to complete research analysis due to technical error.",
                "model_used": "gemini-research",
                "research_timestamp": datetime.utcnow().isoformat(),
                "report_confidence": "Low"
            }

    def generate_legal_report(self, analysis_results: Dict[str, Any], report_focus: str = "comprehensive") -> Dict[str, Any]:
        """
        Generate a comprehensive legal report based on query analysis

        Args:
            analysis_results: Results from analyze_legal_query
            report_focus: Type of report focus (comprehensive, constitutional, criminal, civil, administrative)

        Returns:
            Dictionary containing the legal report
        """
        analysis = analysis_results.get("analysis", {})

        focus_prompts = {
            "comprehensive": """
            Generate a comprehensive legal report covering all aspects of the query.
            Include legal analysis, applicable laws, practical guidance, and recommendations.
            """,
            "constitutional": """
            Focus on constitutional law aspects, fundamental rights, and constitutional remedies.
            Include relevant constitutional provisions and their interpretation.
            """,
            "criminal": """
            Focus on criminal law aspects, offenses, penalties, and criminal procedure.
            Include relevant criminal statutes and procedural requirements.
            """,
            "civil": """
            Focus on civil law aspects, contracts, torts, and civil remedies.
            Include relevant civil statutes and case law precedents.
            """,
            "administrative": """
            Focus on administrative law aspects, government regulations, and administrative procedures.
            Include relevant administrative statutes and regulatory frameworks.
            """
        }

        prompt = f"""
        Based on the following legal query analysis, generate a detailed legal report:

        ANALYSIS RESULTS:
        {json.dumps(analysis, indent=2)}

        REPORT FOCUS: {report_focus}

        {focus_prompts.get(report_focus, focus_prompts["comprehensive"])}

        STRUCTURE THE LEGAL REPORT AS FOLLOWS:

        1. QUERY SUMMARY
           - Original legal question
           - Key issues identified
           - Legal context

        2. APPLICABLE LAW
           - Relevant statutes and regulations
           - Constitutional provisions
           - Case law and precedents

        3. LEGAL ANALYSIS
           - Detailed legal reasoning
           - Rights and obligations
           - Legal implications and consequences

        4. PROCEDURAL GUIDANCE
           - Required steps and procedures
           - Documentation needed
           - Relevant institutions and agencies

        5. PRACTICAL CONSIDERATIONS
           - Implementation challenges
           - Cost and time considerations
           - Alternative approaches

        6. RECOMMENDATIONS
           - Actionable legal recommendations
           - Risk mitigation strategies
           - When to consult legal professionals

        7. ADDITIONAL RESOURCES
           - Further reading and references
           - Related legal topics
           - Support services and organizations

        8. CONCLUSION
           - Summary of legal position
           - Final recommendations
           - Next steps

        Format the report in a professional, accessible manner suitable for non-lawyers while maintaining legal accuracy.
        Include specific references to Kenyan laws, acts, and regulations where applicable.
        """

        try:
            response = self.model.generate_content(prompt)
            report_content = response.text

            return {
                "report": {
                    "title": f"Legal Report - {report_focus.title()} Analysis",
                    "content": report_content,
                    "focus_area": report_focus,
                    "generated_at": datetime.utcnow().isoformat(),
                    "model_used": "gemini-1.5-pro"
                },
                "metadata": {
                    "analysis_timestamp": analysis_results.get("timestamp"),
                    "report_focus": report_focus,
                    "report_length": len(report_content),
                    "sections": self._extract_report_sections(report_content)
                },
                "source_analysis": analysis_results
            }

        except Exception as e:
            logger.error(f"Error generating legal report: {e}")
            return {
                "error": str(e),
                "timestamp": datetime.utcnow().isoformat(),
                "report_focus": report_focus
            }

    def conduct_legal_research(self, legal_topics: List[str], research_questions: List[str]) -> Dict[str, Any]:
        """
        Conduct legal research on specific topics related to Kenya's laws

        Args:
            legal_topics: List of legal topics to research
            research_questions: List of specific research questions

        Returns:
            Dictionary containing legal research findings
        """
        prompt = f"""
        Conduct comprehensive legal research on Kenya's laws based on the following:

        LEGAL TOPICS:
        {json.dumps(legal_topics, indent=2)}

        RESEARCH QUESTIONS:
        {json.dumps(research_questions, indent=2)}

        For each legal topic, provide:
        1. Overview of the legal framework
        2. Key statutes and regulations
        3. Recent developments and amendments
        4. Practical implications for citizens
        5. Common issues and challenges

        Then address each research question with:
        - Legal analysis and interpretation
        - Relevant case law and precedents
        - Practical guidance and procedures
        - Recommendations for legal compliance

        Focus on providing accurate, up-to-date information about Kenyan law.
        Include references to specific laws, acts, and legal sources.
        Provide information that helps users understand and navigate Kenya's legal system.
        """

        try:
            response = self.model.generate_content(prompt)
            research_findings = response.text

            return {
                "legal_research": {
                    "findings": research_findings,
                    "topics_researched": legal_topics,
                    "questions_addressed": research_questions,
                    "generated_at": datetime.utcnow().isoformat(),
                    "model_used": "gemini-1.5-pro"
                },
                "metadata": {
                    "topic_count": len(legal_topics),
                    "question_count": len(research_questions),
                    "findings_length": len(research_findings)
                }
            }

        except Exception as e:
            logger.error(f"Error conducting legal research: {e}")
            return {
                "error": str(e),
                "timestamp": datetime.utcnow().isoformat(),
                "topics": legal_topics,
                "questions": research_questions
            }

    def _fallback_text_analysis(self, query: str, context: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Fallback analysis when JSON parsing fails"""
        context_info = ""
        if context:
            context_info = f"\n\nADDITIONAL CONTEXT:\n{json.dumps(context, indent=2)}"

        prompt = f"""
        Analyze the following legal query about Kenya's laws and provide comprehensive information:

        LEGAL QUERY: {query}{context_info}

        Provide a detailed analysis covering all legal aspects, applicable laws, practical guidance, and recommendations.
        Include specific references to Kenyan laws and legal sources.
        """

        try:
            response = self.model.generate_content(prompt)
            text_analysis = response.text

            return {
                "original_query": query,
                "executive_summary": self._extract_section(text_analysis, "EXECUTIVE SUMMARY") or "Comprehensive legal analysis completed.",
                "detailed_legal_analysis": text_analysis,
                "applicable_laws": self._extract_section(text_analysis, "APPLICABLE LAWS") or "Analysis of relevant Kenyan laws conducted.",
                "practical_guidance": self._extract_section(text_analysis, "PRACTICAL GUIDANCE") or "Legal guidance provided in the analysis.",
                "recommendations": self._extract_section(text_analysis, "RECOMMENDATIONS") or "Legal recommendations included.",
                "model_used": "gemini-research",
                "research_timestamp": datetime.utcnow().isoformat(),
                "report_confidence": "Medium",
                "fallback_mode": True
            }
        except Exception as e:
            return {
                "original_query": query,
                "error": f"Fallback analysis failed: {str(e)}",
                "model_used": "gemini-research",
                "research_timestamp": datetime.utcnow().isoformat(),
                "report_confidence": "Low"
            }

    def generate_pdf_report(self, analysis_results: Dict[str, Any], output_path: str) -> str:
        """
        Generate a PDF document from legal analysis results

        Args:
            analysis_results: Results from analyze_legal_query
            output_path: Path where to save the PDF file

        Returns:
            Path to the generated PDF file
        """
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.units import inch
        except ImportError:
            raise ImportError("reportlab package required for PDF generation. Install with: pip install reportlab")

        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=15,
            textColor=colors.darkblue
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubheading',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=10,
            textColor=colors.darkgreen
        )
        
        normal_style = styles['Normal']
        normal_style.fontSize = 10
        normal_style.leading = 12

        story = []

        # Title Page
        story.append(Paragraph("LEGAL RESEARCH REPORT", title_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("AmaniQuery Legal Intelligence Platform", styles['Heading2']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f"Query: {analysis_results.get('original_query', 'Legal Research Query')}", normal_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Generated: {analysis_results.get('research_timestamp', datetime.utcnow().isoformat())}", normal_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Report Confidence: {analysis_results.get('report_confidence', 'High')}", normal_style))
        story.append(PageBreak())

        # Executive Summary
        if 'executive_summary' in analysis_results:
            story.append(Paragraph("EXECUTIVE SUMMARY", heading_style))
            story.append(Paragraph(analysis_results['executive_summary'], normal_style))
            story.append(Spacer(1, 0.2*inch))

        # Background Context
        if 'background_context' in analysis_results:
            story.append(Paragraph("BACKGROUND CONTEXT", heading_style))
            story.append(Paragraph(analysis_results['background_context'], normal_style))
            story.append(Spacer(1, 0.2*inch))

        # Applicable Laws
        if 'applicable_laws' in analysis_results and isinstance(analysis_results['applicable_laws'], list):
            story.append(Paragraph("APPLICABLE LAWS", heading_style))
            for law in analysis_results['applicable_laws']:
                if isinstance(law, dict):
                    story.append(Paragraph(f"<b>{law.get('law_name', 'Law')}</b>", subheading_style))
                    if 'citation' in law:
                        story.append(Paragraph(f"<i>Citation:</i> {law['citation']}", normal_style))
                    if 'key_provisions' in law and law['key_provisions']:
                        story.append(Paragraph(f"<i>Key Provisions:</i> {', '.join(law['key_provisions'])}", normal_style))
                    if 'interpretation' in law:
                        story.append(Paragraph(f"<i>Application:</i> {law['interpretation']}", normal_style))
                    story.append(Spacer(1, 0.1*inch))
                else:
                    story.append(Paragraph(str(law), normal_style))
            story.append(Spacer(1, 0.2*inch))

        # Legal Analysis
        if 'detailed_legal_analysis' in analysis_results:
            story.append(Paragraph("DETAILED LEGAL ANALYSIS", heading_style))
            story.append(Paragraph(analysis_results['detailed_legal_analysis'], normal_style))
            story.append(Spacer(1, 0.2*inch))

        # Practical Guidance
        if 'practical_guidance' in analysis_results:
            story.append(Paragraph("PRACTICAL GUIDANCE", heading_style))
            guidance = analysis_results['practical_guidance']
            if isinstance(guidance, dict):
                if 'immediate_steps' in guidance and guidance['immediate_steps']:
                    story.append(Paragraph("<b>Immediate Steps:</b>", subheading_style))
                    for step in guidance['immediate_steps']:
                        story.append(Paragraph(f"• {step}", normal_style))
                    story.append(Spacer(1, 0.1*inch))
                
                if 'required_documents' in guidance and guidance['required_documents']:
                    story.append(Paragraph("<b>Required Documents:</b>", subheading_style))
                    for doc in guidance['required_documents']:
                        story.append(Paragraph(f"• {doc}", normal_style))
                    story.append(Spacer(1, 0.1*inch))
                
                if 'relevant_institutions' in guidance and guidance['relevant_institutions']:
                    story.append(Paragraph("<b>Relevant Institutions:</b>", subheading_style))
                    for inst in guidance['relevant_institutions']:
                        story.append(Paragraph(f"• {inst}", normal_style))
                    story.append(Spacer(1, 0.1*inch))
            else:
                story.append(Paragraph(str(guidance), normal_style))
            story.append(Spacer(1, 0.2*inch))

        # Recommendations
        if 'recommendations' in analysis_results and isinstance(analysis_results['recommendations'], list):
            story.append(Paragraph("RECOMMENDATIONS", heading_style))
            for rec in analysis_results['recommendations']:
                if isinstance(rec, dict):
                    priority = rec.get('priority', 'Medium')
                    action = rec.get('action', 'Recommendation')
                    rationale = rec.get('rationale', '')
                    
                    story.append(Paragraph(f"<b>{priority} Priority:</b> {action}", subheading_style))
                    if rationale:
                        story.append(Paragraph(f"<i>Rationale:</i> {rationale}", normal_style))
                    story.append(Spacer(1, 0.1*inch))
                else:
                    story.append(Paragraph(f"• {str(rec)}", normal_style))
            story.append(Spacer(1, 0.2*inch))

        # Disclaimer
        if 'disclaimer' in analysis_results:
            story.append(Paragraph("LEGAL DISCLAIMER", heading_style))
            story.append(Paragraph(analysis_results['disclaimer'], normal_style))
            story.append(Spacer(1, 0.2*inch))

        # Build PDF
        doc.build(story)
        return output_path

    def generate_word_report(self, analysis_results: Dict[str, Any], output_path: str) -> str:
        """
        Generate a Word document from legal analysis results

        Args:
            analysis_results: Results from analyze_legal_query
            output_path: Path where to save the Word file

        Returns:
            Path to the generated Word file
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError("python-docx package required for Word generation. Install with: pip install python-docx")

        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('LEGAL RESEARCH REPORT', 0)
        title.alignment = 1  # Center alignment
        
        # Subtitle
        subtitle = doc.add_heading('AmaniQuery Legal Intelligence Platform', 1)
        subtitle.alignment = 1
        
        # Query and metadata
        doc.add_paragraph(f"Query: {analysis_results.get('original_query', 'Legal Research Query')}")
        doc.add_paragraph(f"Generated: {analysis_results.get('research_timestamp', datetime.utcnow().isoformat())}")
        doc.add_paragraph(f"Report Confidence: {analysis_results.get('report_confidence', 'High')}")
        doc.add_page_break()

        # Executive Summary
        if 'executive_summary' in analysis_results:
            doc.add_heading('EXECUTIVE SUMMARY', 1)
            doc.add_paragraph(analysis_results['executive_summary'])

        # Background Context
        if 'background_context' in analysis_results:
            doc.add_heading('BACKGROUND CONTEXT', 1)
            doc.add_paragraph(analysis_results['background_context'])

        # Applicable Laws
        if 'applicable_laws' in analysis_results and isinstance(analysis_results['applicable_laws'], list):
            doc.add_heading('APPLICABLE LAWS', 1)
            for law in analysis_results['applicable_laws']:
                if isinstance(law, dict):
                    doc.add_heading(law.get('law_name', 'Law'), 2)
                    if 'citation' in law:
                        doc.add_paragraph(f"Citation: {law['citation']}", style='Intense Quote')
                    if 'key_provisions' in law and law['key_provisions']:
                        doc.add_paragraph(f"Key Provisions: {', '.join(law['key_provisions'])}")
                    if 'interpretation' in law:
                        doc.add_paragraph(f"Application: {law['interpretation']}")
                else:
                    doc.add_paragraph(str(law))

        # Legal Analysis
        if 'detailed_legal_analysis' in analysis_results:
            doc.add_heading('DETAILED LEGAL ANALYSIS', 1)
            doc.add_paragraph(analysis_results['detailed_legal_analysis'])

        # Practical Guidance
        if 'practical_guidance' in analysis_results:
            doc.add_heading('PRACTICAL GUIDANCE', 1)
            guidance = analysis_results['practical_guidance']
            if isinstance(guidance, dict):
                if 'immediate_steps' in guidance and guidance['immediate_steps']:
                    doc.add_heading('Immediate Steps', 2)
                    for step in guidance['immediate_steps']:
                        doc.add_paragraph(f"• {step}", style='List Bullet')
                
                if 'required_documents' in guidance and guidance['required_documents']:
                    doc.add_heading('Required Documents', 2)
                    for doc_item in guidance['required_documents']:
                        doc.add_paragraph(f"• {doc_item}", style='List Bullet')
                
                if 'relevant_institutions' in guidance and guidance['relevant_institutions']:
                    doc.add_heading('Relevant Institutions', 2)
                    for inst in guidance['relevant_institutions']:
                        doc.add_paragraph(f"• {inst}", style='List Bullet')
            else:
                doc.add_paragraph(str(guidance))

        # Recommendations
        if 'recommendations' in analysis_results and isinstance(analysis_results['recommendations'], list):
            doc.add_heading('RECOMMENDATIONS', 1)
            for rec in analysis_results['recommendations']:
                if isinstance(rec, dict):
                    priority = rec.get('priority', 'Medium')
                    action = rec.get('action', 'Recommendation')
                    rationale = rec.get('rationale', '')
                    
                    doc.add_heading(f"{priority} Priority: {action}", 2)
                    if rationale:
                        doc.add_paragraph(f"Rationale: {rationale}")
                else:
                    doc.add_paragraph(f"• {str(rec)}", style='List Bullet')

        # Disclaimer
        if 'disclaimer' in analysis_results:
            doc.add_heading('LEGAL DISCLAIMER', 1)
            doc.add_paragraph(analysis_results['disclaimer'])

        # Save document
        doc.save(output_path)
        return output_path

    def _extract_report_sections(self, report_content: str) -> List[str]:
        """Extract section headers from report content"""
        sections = []
        lines = report_content.split('\n')
        for line in lines:
            line = line.strip()
            if line and len(line) < 100:  # Likely a section header
                # Check for common section patterns
                if any(keyword in line.upper() for keyword in [
                    'SUMMARY', 'ANALYSIS', 'GUIDANCE', 'RECOMMENDATIONS', 
                    'CONCLUSION', 'BACKGROUND', 'LAW', 'PROCEDURE'
                ]):
                    sections.append(line)
        return sections

    def _extract_section(self, text: str, section_name: str) -> Optional[str]:
        """Extract a specific section from text content"""
        # Simple text extraction - look for section headers
        lines = text.split('\n')
        in_section = False
        section_content = []
        
        for line in lines:
            line = line.strip()
            if section_name.upper() in line.upper():
                in_section = True
                continue
            elif in_section and line and len(line) > 50:  # Section content
                section_content.append(line)
            elif in_section and line and any(keyword in line.upper() for keyword in [
                'SUMMARY', 'ANALYSIS', 'GUIDANCE', 'RECOMMENDATIONS', 
                'CONCLUSION', 'BACKGROUND', 'LAW', 'PROCEDURE'
            ]):
                # Next section started
                break
        
        return ' '.join(section_content) if section_content else None